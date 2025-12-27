from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

doc = Document()

# HELPERS
def remove_header_spacing(p):
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), "0")
    spacing.set(qn('w:after'), "0")
    spacing.set(qn('w:line'), "240")
    spacing.set(qn('w:lineRule'), "auto")
    pPr.append(spacing)

def table_spacing(p):
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), "360")  # 1.5 line spacing
    spacing.set(qn('w:lineRule'), "auto")
    pPr.append(spacing)

def header(text, size, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    remove_header_spacing(p)

# HEADER
header("FORM ‘A’", 13, True)
header("MEDIATION APPLICATION FORM", 12, True)
header("[REFER RULE 3(1)]", 11, True)
header("Mumbai District Legal Services Authority", 11)
header("City Civil Court, Mumbai", 11)

gap = doc.add_paragraph("")
remove_header_spacing(gap)

# TABLE
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"

table.columns[0].width = Inches(0.6)
table.columns[1].width = Inches(2.2)
table.columns[2].width = Inches(4.2)

# DETAILS OF PARTIES
row = table.rows[0].cells
row[0].merge(row[2])
p = row[0].paragraphs[0]
p.add_run("DETAILS OF PARTIES:").bold = True
table_spacing(p)

# 1 Name of Applicant
r = table.add_row().cells
r[0].paragraphs[0].add_run("1")
r[1].paragraphs[0].add_run("Name of\nApplicant").bold = True
r[2].paragraphs[0].add_run("{{client_name}}")
table_spacing(r[2].paragraphs[0])

# Address and contact details of Applicant
r = table.add_row().cells
r[1].paragraphs[0].add_run("Address and contact details of Applicant").bold = True
r[1].merge(r[2])
table_spacing(r[1].paragraphs[0])

# Address
r = table.add_row().cells
r[1].paragraphs[0].add_run("Address").bold = True
p = r[2].paragraphs[0]
p.add_run("REGISTERED ADDRESS:\n").bold = True
p.add_run("{{branch_address}}\n")
p.add_run("CORRESPONDENCE BRANCH ADDRESS:\n").bold = True
p.add_run("{{branch_address}}")
table_spacing(p)

# Telephone
r = table.add_row().cells
r[1].paragraphs[0].add_run("Telephone No.").bold = True
r[2].paragraphs[0].add_run("{{mobile}}")
table_spacing(r[2].paragraphs[0])

# Mobile
r = table.add_row().cells
r[1].paragraphs[0].add_run("Mobile No.").bold = True

# Email
r = table.add_row().cells
r[1].paragraphs[0].add_run("Email ID").bold = True
email = r[2].paragraphs[0].add_run("info@kslegal.co.in")
email.font.underline = True
table_spacing(r[2].paragraphs[0])

# 2 Opposite Party
r = table.add_row().cells
r[0].paragraphs[0].add_run("2")
p = r[1].paragraphs[0]
p.add_run("Name, Address and Contact details of Opposite Party:").bold = True
r[1].merge(r[2])
table_spacing(p)

# Defendant details
r = table.add_row().cells
p = r[1].paragraphs[0]
p.add_run("Address and contact details of Defendant/s").bold = True
r[1].merge(r[2])
table_spacing(p)

# Name
r = table.add_row().cells
r[1].paragraphs[0].add_run("Name").bold = True
r[2].paragraphs[0].add_run("{{customer_name}}")
table_spacing(r[2].paragraphs[0])

# Address
r = table.add_row().cells
r[1].paragraphs[0].add_run("Address").bold = True
p = r[2].paragraphs[0]
p.add_run("REGISTERED ADDRESS:\n").bold = True
p.add_run("{% if address1 and address1 != \"\" %}{{address1}}{% else %}________________{% endif %}\n")
p.add_run("CORRESPONDENCE ADDRESS:\n").bold = True
p.add_run("{% if address1 and address1 != \"\" %}{{address1}}{% else %}________________{% endif %}")
table_spacing(p)

# Telephone / Mobile / Email
for label in ["Telephone No.", "Mobile No.", "Email ID"]:
    r = table.add_row().cells
    r[1].paragraphs[0].add_run(label).bold = True

# DETAILS OF DISPUTE
r = table.add_row().cells
r[0].merge(r[2])
p = r[0].paragraphs[0]
p.add_run("DETAILS OF DISPUTE:").bold = True
table_spacing(p)

r = table.add_row().cells
r[0].merge(r[2])
p = r[0].paragraphs[0]
p.add_run("THE COMM. COURTS (PRE-INSTITUTION SETTLEMENT) RULES,2018\n").bold = True
p.add_run("Nature of disputes as per section 2(1)(c) of the Commercial Courts Act, 2015 (4 of 2016):")
table_spacing(p)

# SAVE 
doc.save("Mediation_Application_Form.docx")
print("Document generated successfully")
